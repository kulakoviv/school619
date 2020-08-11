from flask import Flask, render_template, request, redirect, send_file, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from docx import Document


app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///rezult.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)


class Prd(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    Name = db.Column(db.String(100), nullable=False)
    Ur = db.Column(db.String(10), nullable=False)
    Kl = db.Column(db.String(5), nullable=False)
    Li = db.Column(db.String(1000))
    Mp = db.Column(db.String(1000))
    Prn = db.Column(db.String(1000))
    Prp = db.Column(db.String(1000))

    def __repr__(self):
        return '<Prd %r>' % self.id


class Redact(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    Name = db.Column(db.String(100), nullable=False)
    Ur = db.Column(db.String(10), nullable=False)
    Kl = db.Column(db.String(5), nullable=False)
    Li = db.Column(db.String(1000))
    Mp = db.Column(db.String(1000))
    Prn = db.Column(db.String(1000))
    Prp = db.Column(db.String(1000))

    def __repr__(self):
        return '<Redact %r>' % self.Name


class Public(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    Name = db.Column(db.String(100), nullable=False)
    Ur = db.Column(db.String(10), nullable=False)
    Kl = db.Column(db.String(5), nullable=False)
    Li = db.Column(db.String(1000))
    Mp = db.Column(db.String(1000))
    Prn = db.Column(db.String(1000))
    Prp = db.Column(db.String(1000))

    def __repr__(self):
        return '<Public %r>' % self.Name


class Res(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    Name = db.Column(db.String(100), nullable=False)
    Ur = db.Column(db.String(10), nullable=False)
    Kl = db.Column(db.String(5), nullable=False)
    Li = db.Column(db.String(1000))
    Mp = db.Column(db.String(1000))
    Prn = db.Column(db.String(1000))
    Prp = db.Column(db.String(1000))

    def __repr__(self):
        return '<Res %r>' % self.Name


@app.route('/')
def index():
    return render_template("index.html")


@app.route('/OOP.docx')
def down():
    return send_file('OOP.docx')


@app.route('/success')
def sc():
    return render_template("success.html")


@app.route('/success_o')
def sc1():
    return render_template("/success_o.html")


@app.route('/create', methods=['POST', 'GET'])
def create():
    if request.method == "POST":
        Pr1 = request.form['Name']
        Ur1 = '(' + request.form['Ur'] + ')'
        Kl1 = request.form['Kl1']
        Kl2 = request.form['Kl2']
        Li1 = ''
        Mp1 = ''
        Prn1 = ''
        Prp1= ''
        for n in range(int(Kl1), int(Kl2)+1):
            article = Prd(Name=Pr1, Ur=Ur1, Kl=str(n), Li=Li1, Mp=Mp1, Prn=Prn1, Prp=Prp1)
            db.session.add(article)
            db.session.commit()
        return redirect('/pr')
    else:
        return render_template("create.html")


@app.route('/kl/<int:id>', methods=['POST', 'GET'])
def red(id):
    article = Prd.query.get(id)
    if request.method == "POST":
        article.Li = request.form['Li']
        article.Mp = request.form['Mp']
        article.Prn = request.form['Prn']
        article.Prp = request.form['Prp']

        try:
            db.session.commit()
            return redirect('')
        except:
            return "Ошибка"
    else:
        return render_template("rab.html", article=article)


@app.route('/adkor/<int:id>', methods=['POST', 'GET'])
def adkor(id):
    article = Prd.query.get(id)
    if request.method == "POST":
        article.Name = request.form['Name']
        article.Ur = request.form['Ur']
        article.Kl = request.form['Kl']
        try:
            db.session.commit()
            return redirect('/pr')
        except:
            return "Ошибка"
    else:
        return render_template("adkor.html", article=article)


@app.route('/adkor1/<int:id>', methods=['POST', 'GET'])
def adkor1(id):
    article = Redact.query.get(id)
    if request.method == "POST":
        article.Name = request.form['Name']
        article.Ur = request.form['Ur']
        article.Kl = request.form['Kl']
        try:
            db.session.commit()
            return redirect('/pr1')
        except:
            return "Ошибка"
    else:
        return render_template("adkor.html", article=article)


@app.route('/pub/<int:id>')
def pub(id):
    article = Prd.query.get(id)
    return render_template("pub.html", article=article)


@app.route('/pub1/<int:id>', methods=['POST', 'GET'])
def pub1(id):
    article = Prd.query.get(id)
    Pr1 = article.Name
    Ur1 = article.Ur
    Kl1 = article.Kl
    Li1 = article.Li
    Mp1 = article.Mp
    Prn1 = article.Prn
    Prp1 = article.Prp
    article1 = Redact(Name=Pr1, Ur=Ur1, Kl=Kl1, Li=Li1, Mp=Mp1, Prn=Prn1, Prp=Prp1)
    db.session.add(article1)
    db.session.delete(article)
    db.session.commit()
    return redirect('/success_o')


@app.route('/pubo/<int:id>', methods=['POST', 'GET'])
def pubo(id):
    article = Redact.query.get(id)
    Pr1 = article.Name
    Ur1 = article.Ur
    Kl1 = article.Kl
    Li1 = article.Li
    Mp1 = article.Mp
    Prn1 = article.Prn
    Prp1 = article.Prp
    article1 = Public(Name=Pr1, Ur=Ur1, Kl=Kl1, Li=Li1, Mp=Mp1, Prn=Prn1, Prp=Prp1)
    db.session.add(article1)
    db.session.delete(article)
    db.session.commit()
    return redirect('/pr1')


@app.route('/addel/<int:id>', methods=['POST', 'GET'])
def adddel(id):
    article = Prd.query.get(id)
    Pr1 = article.Name
    Ur1 = article.Ur
    Kl1 = article.Kl
    Li1 = article.Li
    Mp1 = article.Mp
    Prn1 = article.Prn
    Prp1 = article.Prp
    article1 = Res(Name=Pr1, Ur=Ur1, Kl=Kl1, Li=Li1, Mp=Mp1, Prn=Prn1, Prp=Prp1)
    db.session.add(article1)
    db.session.delete(article)
    db.session.commit()
    return redirect('/pr')


@app.route('/addel1/<int:id>', methods=['POST', 'GET'])
def adddel1(id):
    article = Redact.query.get(id)
    Pr1 = article.Name
    Ur1 = article.Ur
    Kl1 = article.Kl
    Li1 = article.Li
    Mp1 = article.Mp
    Prn1 = article.Prn
    Prp1 = article.Prp
    article1 = Prd(Name=Pr1, Ur=Ur1, Kl=Kl1, Li=Li1, Mp=Mp1, Prn=Prn1, Prp=Prp1)
    db.session.add(article1)
    db.session.delete(article)
    db.session.commit()
    return redirect('/adminsuccess')


@app.route('/oop')
def oop():
    article =[]
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '1').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '2').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '3').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '4').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '5').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '6').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '7').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '8').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '9').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '10').all()))
    article.append(len(Public.query.order_by(Public.Name).filter(Public.Kl == '11').all()))
    a = 0
    for _ in range(11):
        a = a + article[_]
    article.append(len(Public.query.order_by(Public.Name).all()))
    article.append(a)
    return render_template("oop.html", article=article)


@app.route('/pubpub')
def pubpub():
    document = Document()
    articles = Prd.query.order_by(Prd.Name).all()
    r = 'бубу'
    for el in articles:
        if r != el.Name:
            document.add_heading(el.Name + ' ' + el.Kl + ' класс', level=1)
        r = el.Name
        document.add_paragraph(el.Kl + ' класс ' + el.Ur + ' уровень:')
        document.add_paragraph('Личностные:', style='Intense Quote')
        document.add_paragraph(el.Li, style='Intense Quote')
        document.add_paragraph('Метапредметные:', style='Intense Quote')
        document.add_paragraph(el.Mp, style='Intense Quote')
        document.add_paragraph('Ученик научится:', style='Intense Quote')
        document.add_paragraph(el.Prn, style='Intense Quote')
        document.add_paragraph('Ученик получит возможность научиться:', style='Intense Quote')
        document.add_paragraph(el.Prp, style='Intense Quote')
    document.save('OOP.docx')
    return render_template("pubpub.html", articles=articles)


@app.route('/1')
def k1():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '1').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '1').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '1').all()
    return render_template("1.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/2')
def k2():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '2').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '2').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '2').all()
    return render_template("2.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/3')
def k3():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '3').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '3').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '3').all()
    return render_template("3.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/4')
def k4():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '4').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '4').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '4').all()
    return render_template("4.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/5')
def k5():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '5').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '5').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '5').all()
    return render_template("5.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/6')
def k6():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '6').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '6').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '6').all()
    return render_template("6.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/7')
def k7():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '7').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '7').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '7').all()
    return render_template("7.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/8')
def k8():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '8').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '8').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '8').all()
    return render_template("8.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/9')
def k9():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '9').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '9').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '9').all()
    return render_template("9.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/10')
def k10():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '10').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '10').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '10').all()
    return render_template("10.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/11')
def k11():
    articles = Prd.query.order_by(Prd.Name).filter(Prd.Kl == '11').all()
    articles1 = Redact.query.order_by(Redact.Name).filter(Redact.Kl == '11').all()
    articles2 = Public.query.order_by(Public.Name).filter(Public.Kl == '11').all()
    return render_template("11.html", articles=articles, articles1=articles1, articles2=articles2)


@app.route('/pr')
def prr():
    articles = Prd.query.order_by(Prd.Kl).all()
    return render_template("pr.html", articles=articles)


@app.route('/pr1')
def prr1():
    articles = Redact.query.order_by(Redact.Kl).all()
    return render_template("pr1.html", articles=articles)


@app.route('/pro')
def prro():
    articles = Public.query.order_by(Public.Kl).all()
    return render_template("pro.html", articles=articles)


@app.route('/prr')
def prrr():
    articles = Res.query.order_by(Res.Kl).all()
    return render_template("pro.html", articles=articles)


@app.route('/619')
def admin():
    return render_template("/619.html")


@app.route('/adminsuccess')
def adsuc():
    return render_template("/adminsuccess.html")


@app.route('/pr/<int:id>')
def ad1(id):
    article = Prd.query.get(id)
    return render_template("ad.html", article=article)


@app.route('/del/<int:id>')
def del1(id):
    article = Public.query.get(id)
    db.session.delete(article)
    db.session.commit()
    return redirect('/adminsuccess')


@app.route('/pr1/<int:id>')
def ad2(id):
    article = Redact.query.get(id)
    return render_template("ad1.html", article=article)


if __name__ == "__main__":
    app.run(debug=True)